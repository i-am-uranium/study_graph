import io
from dataclasses import dataclass
from enum import StrEnum
from pathlib import Path

import httpx
from app.core.config import get_settings
from docx import Document as DocxDocument
from pypdf import PdfReader, PdfWriter


class DocumentFormat(StrEnum):
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "markdown"
    TEXT = "text"


MIME_TYPES = {
    DocumentFormat.PDF: {"application/pdf"},
    DocumentFormat.DOCX: {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    },
    DocumentFormat.MARKDOWN: {"text/markdown"},
    DocumentFormat.TEXT: {"text/plain"},
}

FILE_EXTENSIONS = {
    DocumentFormat.PDF: {".pdf"},
    DocumentFormat.DOCX: {".docx"},
    DocumentFormat.MARKDOWN: {".md", ".markdown"},
    DocumentFormat.TEXT: {".txt"},
}


def resolve_document_format(content_type: str, suffix: str) -> DocumentFormat:
    normalized_mime = content_type.lower().split(";")[0].strip()
    normalized_ext = suffix.lower()

    # 1. Try to match by MIME Type
    for fmt, mime_set in MIME_TYPES.items():
        if normalized_mime in mime_set:
            return fmt

    # 2. Try to match by File Extension (fallback)
    for fmt, ext_set in FILE_EXTENSIONS.items():
        if normalized_ext in ext_set:
            return fmt

    raise ValueError(f"Unsupported document format: {content_type} ({suffix})")


def split_pdf_to_batches(path: Path, batch_size: int = 10) -> list[tuple[int, bytes]]:
    reader = PdfReader(path)
    num_pages = len(reader.pages)
    batches = []

    for i in range(0, num_pages, batch_size):
        writer = PdfWriter()
        end_page = min(i + batch_size, num_pages)
        for page_num in range(i, end_page):
            writer.add_page(reader.pages[page_num])

        buf = io.BytesIO()
        writer.write(buf)
        batches.append((i, buf.getvalue()))

    return batches


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    metadata: dict


def parse_document(path: Path, content_type: str) -> ParsedDocument:
    fmt = resolve_document_format(content_type, path.suffix)

    match fmt:
        case DocumentFormat.TEXT | DocumentFormat.MARKDOWN:
            return ParsedDocument(
                text=path.read_text(encoding="utf-8", errors="replace"),
                metadata={"source_type": fmt.value},
            )

        case DocumentFormat.PDF:
            settings = get_settings()
            url = settings.unstructured_api_url

            # Split PDF locally into batches of 10 pages to avoid container OOM or timeouts
            batches = split_pdf_to_batches(path, batch_size=10)
            page_to_texts = {}

            for start_page_idx, pdf_bytes in batches:
                files = {
                    "files": (
                        f"batch_{start_page_idx}.pdf",
                        pdf_bytes,
                        "application/pdf",
                    )
                }
                response = httpx.post(url, files=files, timeout=120.0)
                response.raise_for_status()
                elements = response.json()

                for element in elements:
                    text = element.get("text", "").strip()
                    if not text:
                        continue
                    metadata = element.get("metadata", {})

                    # Unstructured page_number is 1-indexed relative to the sub-PDF batch sent.
                    # Adjust it by adding start_page_idx to get the original document page number.
                    sub_page_num = metadata.get("page_number") or 1
                    actual_page_num = start_page_idx + sub_page_num

                    if actual_page_num not in page_to_texts:
                        page_to_texts[actual_page_num] = []
                    page_to_texts[actual_page_num].append(text)

            page_text = []
            sorted_pages = sorted(page_to_texts.keys())
            for page in sorted_pages:
                page_content = "\n".join(page_to_texts[page])
                page_text.append(f"\n\n[Page {page}]\n{page_content}")

            page_count = len(sorted_pages) if sorted_pages else 0

            return ParsedDocument(
                text="\n".join(page_text).strip(),
                metadata={"source_type": fmt.value, "page_count": page_count},
            )

        case DocumentFormat.DOCX:
            doc = DocxDocument(str(path))
            paragraphs = [
                paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
            ]
            return ParsedDocument(
                text="\n".join(paragraphs),
                metadata={"source_type": fmt.value, "paragraph_count": len(paragraphs)},
            )
